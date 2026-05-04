"""
Generate FocusFlow B.Tech report front matter: TOC, List of Tables, List of Figures,
Nomenclature & Abbreviations. Output: ../FOCUSFLOW_Report_TOC_LOT_LOF_Nomenclature.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches, Pt


def set_body_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    f = style.font
    f.name = "Times New Roman"
    f.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)


def add_center_title(doc: Document, text: str, bold: bool = True, size_pt: int = 14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size_pt)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)


def add_leader_line(doc: Document, left: str, page: str, tab_inches: float = 6.25) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(tab_inches), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(left)
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    p.add_run("\t")
    r2 = p.add_run(page)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)


def add_nomenclature_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Symbol / Abbreviation"
    hdr[1].text = "Description / Expansion"
    for c in hdr:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    for i, (abbr, desc) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = abbr
        row[1].text = desc
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "FOCUSFLOW_Report_TOC_LOT_LOF_Nomenclature.docx"

    doc = Document()
    set_body_font(doc)

    # --- Title block (typical inner title page / divider) ---
    doc.add_paragraph()
    add_center_title(doc, "FOCUSFLOW", True, 16)
    add_center_title(
        doc,
        "Execution-First Productivity Platform\n(Mobile Client + REST API + Admin Operations)",
        True,
        12,
    )
    add_center_title(doc, "B.Tech Major Project — Front Matter", False, 12)
    doc.add_paragraph()

    # --- TABLE OF CONTENTS (align with full project report chapters) ---
    add_section_heading(doc, "TABLE OF CONTENTS")
    toc = [
        ("Certificate", "i"),
        ("Declaration", "ii"),
        ("Acknowledgement", "iii"),
        ("Abstract", "iv"),
        ("List of Tables", "v"),
        ("List of Figures", "vi"),
        ("Nomenclature and Abbreviations", "vii"),
        ("1. Introduction", "1"),
        ("1.1 Problem Statement & Motivation", "1"),
        ("1.2 Objectives", "2"),
        ("1.3 Scope of the Project", "2"),
        ("1.4 Organization of the Report", "3"),
        ("2. Literature Review / Related Work", "4"),
        ("3. System Analysis & Requirements", "6"),
        ("3.1 Functional Requirements", "6"),
        ("3.2 Non-Functional Requirements", "7"),
        ("3.3 Use Cases / User Journeys", "8"),
        ("4. System Design", "10"),
        ("4.1 High-Level Architecture (Flutter ↔ NestJS ↔ PostgreSQL)", "10"),
        ("4.2 Database Design (Prisma / ER overview)", "12"),
        ("4.3 API Design (REST /v1, Auth, Notes, Timeline, Planner, AI)", "13"),
        ("4.4 Mobile App Design (Riverpod, go_router, offline-first timeline)", "15"),
        ("4.5 Admin Panel Design (Next.js, role-based admin APIs)", "17"),
        ("5. Implementation Details", "19"),
        ("5.1 Backend (NestJS, Prisma, JWT, logging, migrations)", "19"),
        ("5.2 Mobile Application (Flutter, secure storage, notifications)", "22"),
        ("5.3 Admin Web Console", "24"),
        ("6. Testing & Validation", "26"),
        ("7. Results, Screenshots & Discussion", "28"),
        ("8. Conclusion & Future Work", "30"),
        ("References", "31"),
        ("Appendices (API samples, environment setup)", "32"),
    ]
    for title, page in toc:
        add_leader_line(doc, title, page)

    # --- LIST OF TABLES ---
    add_section_heading(doc, "LIST OF TABLES")
    lots = [
        ("Table 1.1", "Comparison with existing productivity / planner applications", "5"),
        ("Table 3.1", "MVP functional requirements mapping (FocusFlow master spec)", "7"),
        ("Table 3.2", "Non-functional requirements (performance, security, availability)", "8"),
        ("Table 4.1", "Major system modules — mobile, API, database, admin", "11"),
        ("Table 4.2", "Representative REST endpoints under /v1 (auth, me, notes, timeline)", "14"),
        ("Table 5.1", "Technology stack summary (NestJS, Prisma, Flutter, PostgreSQL)", "20"),
        ("Table 6.1", "Test cases / validation matrix (sample)", "27"),
    ]
    for num, caption, page in lots:
        add_leader_line(doc, f"{num}  {caption}", page)

    # --- LIST OF FIGURES ---
    add_section_heading(doc, "LIST OF FIGURES")
    lofs = [
        ("Figure 1.1", "FocusFlow conceptual flow — stuck → next action → focus session", "2"),
        ("Figure 3.1", "Use case diagram — end user vs admin vs system services", "9"),
        ("Figure 4.1", "High-level deployment / component diagram", "11"),
        ("Figure 4.2", "Simplified ER / data model overview (users, notes, timeline, planner)", "12"),
        ("Figure 4.3", "Mobile navigation map (Day 0, shell tabs, focus, add task)", "16"),
        ("Figure 5.1", "Sequence diagram — login / refresh token (JWT) interaction", "21"),
        ("Figure 7.1", "Representative screenshots — Timeline / Now strip", "28"),
        ("Figure 7.2", "Representative screenshots — Deep focus session UI", "29"),
        ("Figure 7.3", "Representative screenshots — Admin panel dashboard (optional)", "29"),
    ]
    for num, caption, page in lofs:
        add_leader_line(doc, f"{num}  {caption}", page)

    # --- NOMENCLATURE ---
    add_section_heading(doc, "NOMENCLATURE AND ABBREVIATIONS")
    doc.add_paragraph(
        "The following symbols, acronyms, and technical terms are used consistently in this "
        "report in the context of the FocusFlow project (mobile + API + admin)."
    )
    nomen_rows: list[tuple[str, str]] = [
        ("API", "Application Programming Interface"),
        ("APK", "Android Application Package (installable build)"),
        ("CI/CD", "Continuous Integration / Continuous Deployment"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("DTO", "Data Transfer Object"),
        ("ER", "Entity–Relationship (database modelling)"),
        ("HTTP / HTTPS", "Hypertext Transfer Protocol (Secure)"),
        ("IDE", "Integrated Development Environment"),
        ("JSON", "JavaScript Object Notation"),
        ("JWT", "JSON Web Token (access / refresh authentication)"),
        ("LAN", "Local Area Network"),
        ("MVP", "Minimum Viable Product"),
        ("ORM", "Object–Relational Mapper (Prisma in this project)"),
        ("OS", "Operating System"),
        ("PWA", "Progressive Web Application"),
        ("REST", "Representational State Transfer (HTTP + JSON style)"),
        ("SDK", "Software Development Kit"),
        ("SQL", "Structured Query Language"),
        ("UI / UX", "User Interface / User Experience"),
        ("URI / URL", "Uniform Resource Identifier / Locator"),
        ("UUID", "Universally Unique Identifier"),
        ("NestJS", "Node.js framework used for the FocusFlow backend"),
        ("Prisma", "Database toolkit and ORM used with PostgreSQL"),
        ("PostgreSQL", "Open-source relational database used by the API"),
        ("Flutter", "Cross-platform UI toolkit (Dart) for the Android client"),
        ("Riverpod", "State management library used in the Flutter app"),
        ("Dio", "HTTP client library for Flutter (REST calls, interceptors)"),
        ("Next.js", "React framework used for the admin web panel"),
        ("OTP / 2FA", "One-Time Password / Two-Factor Authentication (if applicable)"),
        ("RBAC", "Role-Based Access Control (admin roles)"),
        ("TWA", "Trusted Web Activity (web-to-store packaging pattern)"),
        ("UTC", "Coordinated Universal Time (server date bucketing)"),
        ("FF", "FocusFlow (internal shorthand in specs / keys)"),
    ]
    add_nomenclature_table(doc, nomen_rows)

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
