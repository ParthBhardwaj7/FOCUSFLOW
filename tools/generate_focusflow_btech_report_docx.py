# -*- coding: utf-8 -*-
"""
Generate CEC-CGC style B.Tech project report DOCX for FocusFlow.
Margins per college format: left 3.5cm, top 2.5cm, right 1.25cm, bottom 1.25cm.
Body: Times New Roman 12pt, 1.5 line spacing (project report spec).
Output: FocusFlow_BTech_Full_Project_Report.docx in repository root.
"""
from __future__ import annotations

import sys
from pathlib import Path

_REPO_ROOT = Path(__file__).resolve().parents[1]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from tools.report_expansion_chunks import ALL_EXPANSION_PARAGRAPHS


def set_section_margins(section) -> None:
    section.left_margin = Cm(3.5)
    section.top_margin = Cm(2.5)
    section.right_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)


def configure_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_title(document: Document, text: str, size_pt: int = 14, bold: bool = True) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)


def add_subheading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


def add_body(document: Document, text: str, first_line_indent: bool = False) -> None:
    p = document.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def page_break(document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def main() -> None:
    document = Document()
    for section in document.sections:
        set_section_margins(section)
    configure_normal_style(document)

    # ----- Cover / inner title -----
    document.add_paragraph()
    add_title(document, "Chandigarh Engineering College", 12, True)
    add_title(document, "Chandigarh Group of Colleges", 12, True)
    add_title(document, "Landran, Mohali — 140 307", 12, False)
    document.add_paragraph()
    add_title(document, "FOCUSFLOW", 24, True)
    add_title(document, "A Cross-Platform Execution Engine for Attention,\nPlanning, and Guided Focus", 14, True)
    document.add_paragraph()
    add_title(document, "B.Tech Project Report", 14, True)
    add_title(document, "Department of Computer Science & Engineering", 12, False)
    document.add_paragraph()
    add_title(document, "Submitted by", 11, True)
    add_title(document, "[Student Name]    PTU Roll No. [Roll Number]", 12, False)
    add_title(document, "[Branch]    Batch [Year]", 12, False)
    add_title(document, "E-mail: [email]    Mobile: [phone]", 12, False)
    document.add_paragraph()
    add_title(document, "Faculty Guide: [Name, Designation]", 12, False)
    page_break(document)

    # ----- Certificate -----
    add_heading(document, "CERTIFICATE")
    add_body(
        document,
        "I hereby certify that the work presented in this B.Tech Project Report entitled "
        "“FocusFlow — A Cross-Platform Execution Engine for Attention, Planning, and Guided Focus”, "
        "in partial fulfilment of the requirements for the award of the Bachelor of Technology in "
        "Computer Science & Engineering, submitted to the Department of Computer Science & Engineering, "
        "CEC-CGC Landran, Mohali, Punjab, is an authentic record of my own work carried out during "
        "the 6th semester under the supervision of [Supervisor Name, Designation].",
    )
    add_body(
        document,
        "The matter presented in this project report has not been submitted by me for the award of "
        "any other degree elsewhere.",
    )
    document.add_paragraph()
    add_body(document, "Signature of Student(s) ______________________    Date: ____________")
    document.add_paragraph()
    add_body(
        document,
        "This is to certify that the above statement made by the student(s) is correct to the best "
        "of my knowledge.",
    )
    add_body(document, "Signature of Supervisor ______________________    Date: ____________")
    add_body(document, "Head of Department, Computer Science & Engineering ______________________")
    page_break(document)

    # ----- Acknowledgement -----
    add_heading(document, "ACKNOWLEDGEMENT")
    add_body(
        document,
        "I would like to record my sincere gratitude to the Head of the Department of Computer "
        "Science & Engineering, CEC-CGC Landran, Mohali, India, for administrative support and for "
        "providing access to laboratory and library facilities.",
        True,
    )
    add_body(
        document,
        "I express my deep appreciation to my project supervisor for continuous guidance, "
        "constructive feedback, and encouragement throughout the development of FocusFlow.",
        True,
    )
    add_body(
        document,
        "I also thank faculty members and peers who reviewed interim demonstrations and offered "
        "practical suggestions that improved the quality of the engineering outcomes.",
        True,
    )
    add_body(document, "Signature of Student ______________________    Name: [Student Name]")
    page_break(document)

    # ----- Abstract -----
    add_heading(document, "ABSTRACT")
    abstract = (
        "Modern knowledge workers and students routinely struggle with procrastination, context "
        "switching, and cognitive overload when using conventional task managers that emphasise "
        "cataloguing work rather than executing it. FocusFlow addresses this gap by positioning "
        "itself as an execution engine: the product minimises the time from hesitation to an "
        "active focus session through a command-first home surface, compassionate recovery flows "
        "after missed blocks, and tight integration between planning (timeline), capture (inbox "
        "notes), and deep work (timer, audio, and session outcomes). The implementation follows "
        "a three-tier architecture: a Flutter mobile client for Android (Riverpod, go_router, "
        "Dio, secure token storage, local notifications, and audio services), a NestJS 11 REST "
        "API versioned under /v1 with PostgreSQL persistence via Prisma 6, and an optional Next.js "
        "administration console for operational governance (feature flags, content, audit, and "
        "user support workflows). Security is treated as a first-class concern: environment "
        "validation with Zod, Helmet, throttling, structured JSON errors, database-backed refresh "
        "token rotation, role-based admin access, and maintenance hooks for remote configuration. "
        "The domain model spans users, tasks, timeline slots, focus sessions, planner day "
        "snapshots for offline-first mirroring, notes (including optional voice attachments), AI "
        "coach logs and suggestions, client error reporting, and push notification infrastructure. "
        "This report documents requirements, literature-informed design rationale, methodology, "
        "system architecture, module-level behaviour, testing posture, observed outcomes from "
        "development builds, limitations near timezone boundaries, and future scope including "
        "expanded analytics, subscription billing, and richer on-device intelligence. The narrative "
        "aligns with institutional expectations for a full-length technical report while mapping "
        "each chapter to verifiable artefacts in the repository (backend README, mobile README, "
        "Prisma schema, and master product specification)."
    )
    add_body(document, abstract, True)
    page_break(document)

    # ----- Table of contents (manual) -----
    add_heading(document, "TABLE OF CONTENTS")
    toc_lines = [
        "Certificate ........................................................ i",
        "Acknowledgement .................................................... ii",
        "Abstract ........................................................... iii",
        "Table of Contents .................................................. iv",
        "List of Tables ..................................................... v",
        "List of Figures .................................................... vi",
        "Abbreviations and Nomenclature ...................................... vii",
        "Chapter 1 Introduction ............................................... 1",
        "Chapter 2 Literature Survey ........................................ 6",
        "Chapter 3 Present Work .............................................. 12",
        "Chapter 4 Results and Discussion .................................... 24",
        "Chapter 5 Conclusion and Future Scope ............................... 28",
        "References ........................................................ 30",
        "Appendices ........................................................ 31",
    ]
    for line in toc_lines:
        add_body(document, line)
    page_break(document)

    add_heading(document, "LIST OF TABLES")
    add_body(document, "Table 3.1 — Representative API surface (versioned /v1) ................ 14")
    add_body(document, "Table 3.2 — Core relational entities (PostgreSQL via Prisma) .......... 18")
    page_break(document)

    add_heading(document, "LIST OF FIGURES")
    add_body(document, "Figure 3.1 — Logical deployment view (mobile, API, database) ....... 13")
    add_body(document, "Figure 3.2 — Mobile navigation shell (tabs and deep routes) ........... 16")
    page_break(document)

    add_heading(document, "ABBREVIATIONS AND NOMENCLATURE")
    add_body(document, "API — Application Programming Interface; CRUD — Create, Read, Update, Delete.")
    add_body(document, "JWT — JSON Web Token; ORM — Object-Relational Mapping; MVP — Minimum Viable Product.")
    add_body(document, "REST — Representational State Transfer; UTC — Coordinated Universal Time.")
    add_body(document, "MIT — Most Important Task (priority flag on tasks in FocusFlow).")
    page_break(document)

    # ----- Chapter 1 -----
    add_heading(document, "CHAPTER 1 — INTRODUCTION")
    add_subheading(document, "1.1 Background and Motivation")
    add_body(
        document,
        "Digital productivity tools have proliferated across mobile and web ecosystems, yet empirical "
        "observations and product critiques consistently highlight a structural mismatch: users "
        "accumulate lists, labels, and dashboards while still deferring the actual unit of progress, "
        "which is sustained attention on a single next action. Cognitive science literature "
        "underscores that excessive choice and elaborate setup rituals increase activation energy, "
        "particularly for populations sensitive to shame after missed intentions. FocusFlow is "
        "motivated by a north-star metric — reducing seconds from a stuck state to a running timer — "
        "and by behavioural principles documented in the project’s master specification: minimise "
        "early cognitive load, avoid guilt amplification after failure, and preserve a credible "
        "path to closure through small wins.",
        True,
    )
    add_body(
        document,
        "The engineering response is not merely motivational copy but a coherent software architecture "
        "that couples a calm, command-first mobile experience with a robust, auditable server capable "
        "of identity, persistence, AI-assisted coaching, and remote operations. The repository "
        "implements this vision using mainstream, industry-proven stacks that remain teachable in "
        "an undergraduate curriculum: Dart/Flutter on the client, TypeScript/NestJS on the server, "
        "PostgreSQL as the system of record, and Prisma as the schema-first data access layer.",
    )

    add_subheading(document, "1.2 Problem Statement")
    add_body(
        document,
        "Students and professionals need a system that answers “what should I do now?” without "
        "forcing them through multi-step configuration. They also need continuity across sessions "
        "(tokens, cached profile), resilience when offline (local planner mirror), and fair handling "
        "of partial failure (network timeouts, revoked refresh tokens). Administrators require "
        "visibility into feature flags, content, and error fingerprints without compromising end-user "
        "privacy or availability.",
        True,
    )

    add_subheading(document, "1.3 Objectives")
    objs = [
        "Deliver a mobile execution loop: Day 0 onboarding, timeline/week strip, focus timer with "
        "audio scaffolding, inbox capture, and recovery-oriented reset-day flows.",
        "Provide a secure REST API with explicit validation, consistent error envelopes, rate "
        "limiting, and database migrations suitable for team development and controlled rollout.",
        "Model productivity domain entities (tasks, slots, sessions, notes, planner snapshots) with "
        "clear indices and ownership rules.",
        "Integrate optional AI coaching with logging for accountability and future quality analysis.",
        "Offer an admin surface for governance (users, flags, notifications, audit trail).",
    ]
    for o in objs:
        add_body(document, "• " + o)

    add_subheading(document, "1.4 Scope and Delimitations")
    add_body(
        document,
        "The present work centres on an Android-first Flutter client wired to a self-hosted Nest "
        "API. The report describes UTC versus local-day semantics where the backend interprets "
        "calendar-day queries; this is a known delimitation for users near midnight boundaries. "
        "Payments and production-scale observability are outlined as future scope rather than fully "
        "implemented product features.",
        True,
    )

    add_subheading(document, "1.5 Organisation of the Report")
    add_body(
        document,
        "Chapter 2 reviews related products and technical foundations. Chapter 3 details analysis, "
        "design, and implementation across mobile, server, database, and admin modules. Chapter 4 "
        "summarises outcomes, risks, and discussion. Chapter 5 concludes with future enhancements.",
        True,
    )
    page_break(document)

    # ----- Chapter 2 -----
    add_heading(document, "CHAPTER 2 — LITERATURE SURVEY")
    add_subheading(document, "2.1 Productivity Applications and User Experience")
    add_body(
        document,
        "Contemporary task managers (Todoist, Microsoft To Do, Apple Reminders, Notion databases) "
        "excel at capture and categorisation yet often surface multiple competing calls-to-action. "
        "FocusFlow differentiates through a command-first “Now” philosophy aligned with research on "
        "habit loops: cue, routine, reward. The literature on compassionate design cautions against "
        "interfaces that moralise missed tasks; FocusFlow encodes recovery affordances such as "
        "compressing the remainder of a day or emphasising MIT-only views.",
        True,
    )

    add_subheading(document, "2.2 Mobile Engineering Patterns")
    add_body(
        document,
        "Flutter enables declarative UI with a single codebase and strong performance on Android. "
        "Riverpod supports testable dependency injection and reactive state. go_router yields typed "
        "navigation graphs including authenticated shells. Secure storage for tokens follows OWASP "
        "mobile guidance. Local notifications and audio services interact with Android foreground "
        "service policies where long-running focus audio is concerned.",
    )

    add_subheading(document, "2.3 Service-Oriented Back Ends")
    add_body(
        document,
        "NestJS modularises cross-cutting concerns (guards, interceptors, pipes) and aligns with "
        "enterprise TypeScript practices. Prisma migrations provide auditable evolution of relational "
        "schemas. PostgreSQL offers transactional integrity for token rotation and timeline writes.",
    )

    add_subheading(document, "2.4 AI Assistance and Governance")
    add_body(
        document,
        "Large language model integrations require logging, rate limits, and content safety "
        "considerations. FocusFlow stores coach exchanges for quality review while keeping PII "
        "handling explicit. Admin tooling for suggestions and templates operationalises responsible "
        "iteration without shipping hard-coded prompt strings exclusively inside mobile binaries.",
    )
    page_break(document)

    # ----- Chapter 3 -----
    add_heading(document, "CHAPTER 3 — PRESENT WORK")
    add_subheading(document, "3.1 Requirements Analysis")
    add_body(
        document,
        "Functional requirements include registration and login, refresh rotation, profile patch, "
        "task CRUD by calendar day, timeline CRUD with ISO8601 instants, focus session start and "
        "completion outcomes, notes with tags and optional voice audio, planner cloud sync snapshots, "
        "AI chat endpoints, client error reporting, and admin-only maintenance configuration. "
        "Non-functional requirements include sub-second median health checks, structured logging "
        "with Pino, request identifiers on responses, CORS allow-lists, throttling defaults, and "
        "readiness probes that verify database connectivity.",
        True,
    )

    add_subheading(document, "3.2 Methodology")
    add_body(
        document,
        "An iterative methodology was adopted: schema-first modelling in Prisma, contract-first REST "
        "routes documented in backend README tables, parallel mobile feature tracks (shell, "
        "timeline, focus, inbox), and incremental hardening (guards, maintenance middleware, audit). "
        "Local Docker Compose supports reproducible PostgreSQL instances for each developer machine.",
    )

    add_subheading(document, "3.3 System Architecture")
    add_body(
        document,
        "Figure 3.1 (logical). Flutter client ↔ HTTPS JSON ↔ Nest API ↔ Prisma ↔ PostgreSQL. "
        "Optional Next.js admin authenticates against admin JWT flows. Mobile uses Dio interceptors "
        "for refresh-on-401. Timeline planner remains local-first with cloud snapshots for "
        "multi-device resilience documented in schema comments.",
    )

    add_subheading(document, "3.4 Backend Modules (NestJS)")
    modules = [
        "Auth: register/login/refresh/logout with hashed refresh tokens.",
        "Users/me: profile retrieval and onboarding completion timestamps.",
        "Tasks: per-day scheduling, MIT flag, archival and completion timestamps.",
        "Timeline: ordered slots with statuses UPCOMING/ACTIVE/DONE/MISSED/SKIPPED.",
        "Focus sessions: planned duration, optional task linkage, JSON subtask snapshots.",
        "Notes: rich text body, comma-separated tags, optional audioKey for uploads.",
        "Planner: upsert of per-day JSON mirrors for offline-first UX.",
        "AI: chat DTO validation, LLM service abstraction, logging to AiCoachLog.",
        "Client hooks: public config, mobile error reports, flags, push registration stubs.",
        "Admin: dashboard metrics, user moderation, feature flags, notifications, categories, "
        "sounds, audit logs, error triage.",
        "Health: /v1/health liveness and /v1/ready database probe.",
    ]
    for m in modules:
        add_body(document, "• " + m)

    add_subheading(document, "3.5 Mobile Application (Flutter)")
    add_body(
        document,
        "Navigation covers splash, authentication, Day 0 onboarding, a tabbed shell (Inbox, "
        "Timeline/Now, AI, Settings), add-task sheet, deep focus pages, note editor, coach context "
        "settings, demographics and focus profile preferences, notification bootstrap, and daily "
        "nudge schedulers. State is coordinated with Riverpod providers; timeline data is persisted "
        "locally while selected calendar semantics are documented for UTC bridging.",
    )

    add_subheading(document, "3.6 Database Design Highlights")
    add_body(
        document,
        "Table 3.2 (illustrative). User anchors one-to-many relations to tasks, slots, sessions, "
        "notes, planner snapshots, memories, devices, and logs. Indices align with query paths such "
        "as (userId, scheduledOn) for tasks and (userId, startsAt) for slots. Error logs capture "
        "fingerprinted occurrences for deduplication in admin triage.",
    )

    add_subheading(document, "3.7 Security and Compliance Considerations")
    add_body(
        document,
        "JWT access secrets meet minimum length checks at boot. Refresh tokens are stored hashed. "
        "Admin routes compose JwtAuthGuard with role guards. Maintenance middleware can short-circuit "
        "requests with structured maintenance payloads for client UX. CORS origins are explicit "
        "comma-separated lists — no wildcard in production configurations.",
    )

    add_subheading(document, "3.8 Administration Console")
    add_body(
        document,
        "The Next.js admin panel provides authenticated dashboards for operational staff: viewing "
        "errors, toggling flags, sending push campaigns, managing taxonomy for sounds/categories, "
        "and reviewing audit trails. This separates product engineering from live configuration, "
        "reducing time-to-mitigate for incidents.",
    )

    add_subheading(document, "3.9 API Surface (Representative)")
    add_body(
        document,
        "Table 3.1. Public: POST /v1/auth/register, POST /v1/auth/login, POST /v1/auth/refresh, "
        "POST /v1/auth/logout. Authenticated: GET/PATCH /v1/me; GET/POST/PATCH/DELETE /v1/tasks; "
        "GET/POST/PATCH/DELETE /v1/timeline; POST/PATCH /v1/focus-sessions; notes and planner routes "
        "as implemented in repository branches; health/ready for probes.",
    )

    add_subheading(document, "3.10 Testing Strategy")
    add_body(
        document,
        "Backend ships e2e smoke with mocked Prisma for CI friendliness; developers validate "
        "/v1/ready against real Postgres locally. Mobile testing relies on flutter analyse, widget "
        "and integration tests where present, and manual device matrices (emulator, USB reverse, "
        "LAN IP) for networking edge cases.",
    )

    add_heading(document, "SUPPLEMENTARY TECHNICAL MONOGRAPH")
    add_body(
        document,
        "The following sections consolidate feasibility analysis, extended software requirements "
        "commentary, methodology elaboration, risk registers, literature bridges, and operational "
        "governance notes. Together they provide the depth expected of a full-length B.Tech project "
        "report while remaining anchored to the FocusFlow repository and master specification.",
        True,
    )
    for block in ALL_EXPANSION_PARAGRAPHS:
        add_body(document, block, True)

    add_subheading(document, "Appendix table — representative REST catalogue (v1)")
    api_lines = [
        "POST /v1/auth/register — body {email,password}; returns tokens and user profile.",
        "POST /v1/auth/login — body {email,password}; same response family as register.",
        "POST /v1/auth/refresh — body {refreshToken}; rotates refresh token row.",
        "POST /v1/auth/logout — body {refreshToken}; revokes presented refresh token.",
        "GET /v1/me — returns current user without password hash.",
        "PATCH /v1/me — partial update for onboardingCompletedAt, timeZone, and related profile fields.",
        "GET /v1/tasks?on=YYYY-MM-DD — lists tasks for UTC calendar bucket MVP.",
        "POST /v1/tasks — creates title, optional notes, scheduledOn, sortOrder, isMit.",
        "PATCH /v1/tasks/:id — owner-only partial updates.",
        "DELETE /v1/tasks/:id — owner-only deletion.",
        "POST /v1/focus-sessions — starts PENDING session with plannedDurationSec and optional taskId.",
        "PATCH /v1/focus-sessions/:id — completes or skips; sets endedAt.",
        "GET /v1/timeline?on=YYYY-MM-DD — lists slots intersecting UTC day window, time ordered.",
        "POST /v1/timeline — creates slot with ISO8601 startsAt/endsAt and optional metadata.",
        "PATCH /v1/timeline/:id — owner-only partial updates.",
        "DELETE /v1/timeline/:id — owner-only removal.",
        "GET /v1/health — liveness probe for process up.",
        "GET /v1/ready — readiness probe executing SELECT 1 style database check.",
    ]
    for line in api_lines:
        add_body(document, line)

    add_subheading(document, "Appendix table — core environment variables (backend)")
    env_lines = [
        "NODE_ENV — development | production | test.",
        "PORT — HTTP listen port (default 3000).",
        "LOG_LEVEL — Pino verbosity control.",
        "DATABASE_URL — required PostgreSQL connection string for Prisma.",
        "SHADOW_DATABASE_URL — optional for hosts requiring shadow DB during migrate.",
        "JWT_ACCESS_SECRET — required secret ≥32 chars for signing access JWTs.",
        "JWT_REFRESH_SECRET — optional legacy key; refresh rows are database-backed.",
        "JWT_ACCESS_EXPIRES_IN — default 15m unless overridden.",
        "JWT_REFRESH_EXPIRES_IN — default 7d unless overridden.",
        "CORS_ORIGINS — required comma-separated allow-list without spaces.",
        "COOKIE_DOMAIN / COOKIE_SECURE — optional cookie hardening for future refresh-cookie flows.",
        "THROTTLE_TTL / THROTTLE_LIMIT — global rate limit knobs; health routes skip throttle.",
        "SENTRY_DSN — optional crash reporting integration.",
    ]
    for line in env_lines:
        add_body(document, line)

    page_break(document)

    # ----- Chapter 4 -----
    add_heading(document, "CHAPTER 4 — RESULTS AND DISCUSSION")
    add_body(
        document,
        "The integrated system demonstrates end-to-end account creation, tokenised session reuse, "
        "task and timeline persistence, and focus session lifecycle transitions from PENDING to "
        "COMPLETED or SKIPPED. Structured JSON errors improve mobile snackbar clarity versus opaque "
        "failures. Admin workflows reduce operational blind spots by centralising error fingerprints "
        "and feature toggles.",
        True,
    )
    add_body(
        document,
        "Discussion: timezone alignment between local planner selection and UTC bucketed queries "
        "remains a teaching moment for specification testing; mitigation paths include documenting "
        "user-facing semantics, server-side TZ-aware windows in future versions, or explicit "
        "offset-aware DTO fields. Audio and notification reliability depend on OEM battery policies; "
        "native foreground services are the industry pattern for long focus audio.",
    )
    add_body(
        document,
        "Performance-wise, indexed lookups keep day-scoped queries responsive for MVP data volumes. "
        "Throttling protects brute-force surfaces on authentication while health routes remain "
        "excluded to support orchestrators.",
    )
    page_break(document)

    # ----- Chapter 5 -----
    add_heading(document, "CHAPTER 5 — CONCLUSION AND FUTURE SCOPE")
    add_body(
        document,
        "FocusFlow successfully concretises an execution-oriented product thesis in code: a Flutter "
        "client, a NestJS API, a normalised PostgreSQL schema, and governance tooling. The project "
        "demonstrates full-stack literacy, secure defaults, and maintainable module boundaries.",
        True,
    )
    add_body(
        document,
        "Future scope includes subscription management (e.g., RevenueCat), richer analytics cohorts "
        "matching master-spec events, improved timezone models, encrypted backup exports, expanded "
        "widget surfaces, and deeper on-device ML for ranking the next block without round-trips.",
    )
    page_break(document)

    add_heading(document, "REFERENCES")
    refs = [
        "Flutter Documentation — Offline-first application architecture patterns.",
        "NestJS Documentation — Fundamentals, security, and testing guides.",
        "Prisma Documentation — Schema modelling and migrations.",
        "PostgreSQL Documentation — Relational integrity and indexing.",
        "OWASP Mobile Security — Secure local storage and transport best practices.",
        "FocusFlow repository internal specifications: FOCUSFLOW_MASTER.md, backend/README.md, mobile/README.md.",
    ]
    for r in refs:
        add_body(document, "• " + r)

    page_break(document)
    add_heading(document, "APPENDIX A — SAMPLE ENVIRONMENT VARIABLES (BACKEND)")
    add_body(
        document,
        "DATABASE_URL, JWT_ACCESS_SECRET (≥32 chars), CORS_ORIGINS, optional SENTRY_DSN, throttle "
        "tuning keys, LOG_LEVEL — validated at boot via Zod schema in env.validation.ts.",
    )
    add_heading(document, "APPENDIX B — ETHICAL NOTE")
    add_body(
        document,
        "AI features must be deployed with transparent privacy disclosures, opt-in where required by "
        "policy, and administrative review of prompts and logged interactions to prevent misuse.",
    )

    out_path = "FocusFlow_BTech_Full_Project_Report.docx"
    document.save(out_path)
    print("Wrote", out_path)


if __name__ == "__main__":
    main()
